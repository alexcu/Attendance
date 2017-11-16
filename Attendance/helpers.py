import os

import pandas
from docx import Document
from pandas import ExcelFile
from pulp import LpProblem, LpMinimize, lpSum, LpVariable, LpStatus, LpInteger, LpBinary

from Attendance import app, db, executor
from Attendance.models import *
import Attendance.models
from Attendance.forms import AddTimetableForm

#TIMETABLE CODE
def runtimetable2(STUDENTS, SUBJECTS, TIMES, day, DAYS, TEACHERS, SUBJECTMAPPING, REPEATS, TEACHERMAPPING,
                  TUTORAVAILABILITY, maxclasssize, minclasssize, nrooms):
    '''
    Run the timetabling process and input into the database.

    This process calls the CBCSolver using the PuLP package and then adds the classes to the database.


    :param STUDENTS: should be an array of student names
    :param SUBJECTS: should be an array of subject codes
    :param TIMES: an array of strings representing possible timeslots
    :param day:
    :param DAYS: the days corresponding to the timeslots above
    :param TEACHERS: an array of the names of the tutors
    :param SUBJECTMAPPING: This is a dictionary representing the subjects
                            each tutor is taking
    :param REPEATS: A dictionary of how many repeats each subject has
    :param TEACHERMAPPING: A dictionary of what subject each tutor teachers
    :param TUTORAVAILABILITY:
    :param maxclasssize: An integer representing the maximum class size
    :param minclasssize: An integer representing the minimum class size
    :param nrooms: An integer representing the max allowable concurrent classes
    :return: A string representing model status.
    '''
    print("Running solver")
    model = LpProblem('Timetabling', LpMinimize)
    # Create Variables
    print("Creating Variables")
    app.logger.info('Assignment Variables')
    assign_vars = LpVariable.dicts("StudentVariables",
                                   [(i, j, k, m) for m in TEACHERS for j in TEACHERMAPPING[m] for i in SUBJECTMAPPING[j]
                                    for k in TIMES], 0, 1, LpBinary)
    app.logger.info('Subject Variables')
    subject_vars = LpVariable.dicts("SubjectVariables",
                                    [(j, k, m) for m in TEACHERS for j in TEACHERMAPPING[m] for k in TIMES], 0, 1,
                                    LpBinary)

    # c
    app.logger.info('9:30 classes')
    num930classes = LpVariable.dicts("930Classes", [(i) for i in TIMES], lowBound=0, cat=LpInteger)
    # w
    app.logger.info('Days for teachers')
    daysforteachers = LpVariable.dicts("numdaysforteachers", [(i, j) for i in TEACHERS for j in range(len(DAYS))], 0, 1,
                                       LpBinary)
    # p
    daysforteacherssum = LpVariable.dicts("numdaysforteacherssum", [(i) for i in TEACHERS], 0, cat=LpInteger)
    # variables for student clashes
    studenttime = LpVariable.dicts("StudentTime", [(i, j) for i in STUDENTS for j in TIMES], lowBound=0, upBound=1,
                                   cat=LpBinary)
    studentsum = LpVariable.dicts("StudentSum", [(i) for i in STUDENTS], 0, cat=LpInteger)

    # Count the days that a teacher is rostered on. Make it bigger than a small number times the sum
    # for that particular day.
    for m in TEACHERS:
        app.logger.info('Counting Teachers for ' + m)
        for d in range(len(day)):
            model += daysforteachers[(m, d)] >= 0.1 * lpSum(
                subject_vars[(j, k, m)] for j in TEACHERMAPPING[m] for k in DAYS[day[d]])
            model += daysforteachers[(m, d)] <= lpSum(
                subject_vars[(j, k, m)] for j in TEACHERMAPPING[m] for k in DAYS[day[d]])
    for m in TEACHERS:
        model += daysforteacherssum[(m)] == lpSum(daysforteachers[(m, d)] for d in range(len(day)))

    print("Constraining tutor availability")
    # This bit of code puts in the constraints for the tutor availability.
    # It reads in the 0-1 matrix of tutor availability and constrains that no classes
    # can be scheduled when a tutor is not available.
    # The last column of the availabilities is the tutor identifying number, hence why we have
    # used a somewhat convoluted idea down here.
    for m in TEACHERS:
        for k in TIMES:
            if k not in TUTORAVAILABILITY[m]:
                model += lpSum(subject_vars[(j, k, m)] for j in TEACHERMAPPING[m]) == 0

    # Constraints on subjects for each students
    print("Constraining student subjects")
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            for i in SUBJECTMAPPING[j]:
                model += lpSum(assign_vars[(i, j, k, m)] for k in TIMES) == 1

    # This code means that students cannot attend a tute when a tute is not running
    # But can not attend a tute if they attend a repeat.
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            for i in SUBJECTMAPPING[j]:
                for k in TIMES:
                    model += assign_vars[(i, j, k, m)] <= subject_vars[(j, k, m)]

    # Constraints on which tutor can take each class
    # This goes through each list and either constrains it to 1 or 0 depending if
    # the teacher needs to teach that particular class.
    print("Constraining tutor classes")
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            model += lpSum(subject_vars[(j, k, m)] for k in TIMES) == REPEATS[j]

    # General Constraints on Rooms etc.
    print("Constraining times")
    # For each time cannot exceed number of rooms
    for k in TIMES:
        model += lpSum(subject_vars[(j, k, m)] for m in TEACHERS for j in TEACHERMAPPING[m]) <= nrooms

    # Teachers can only teach one class at a time
    for k in TIMES:
        for m in TEACHERS:
            model += lpSum(subject_vars[(j, k, m)] for j in TEACHERMAPPING[m]) <= 1
    print("Constraint: Minimize student clashes")
    # STUDENT CLASHES
    for i in STUDENTS:
        for k in TIMES:
            model += studenttime[(i, k)] <= lpSum(
                assign_vars[(i, j, k, m)] for m in TEACHERS for j in TEACHERMAPPING[m] if i in SUBJECTMAPPING[j]) / 2
            model += studenttime[(i, k)] >= 0.3 * (0.5 * lpSum(
                assign_vars[(i, j, k, m)] for m in TEACHERS for j in TEACHERMAPPING[m] if i in SUBJECTMAPPING[j]) - 0.5)
    for i in STUDENTS:
        model += studentsum[(i)] == lpSum(studenttime[(i, k)] for k in TIMES)

    # This minimizes the number of 9:30 classes.
    for i in TIMES:
        if i.find('21:30') != -1:
            model += num930classes[(i)] == lpSum(subject_vars[(j, i, m)] for m in TEACHERS for j in TEACHERMAPPING[m])

        else:
            model += num930classes[(i)] == 0

    print("Setting objective function")

    # Class size constraint
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            for k in TIMES:
                model += lpSum(assign_vars[(i, j, k, m)] for i in SUBJECTMAPPING[j]) >= minclasssize * subject_vars[
                    (j, k, m)]
                model += lpSum(assign_vars[(i, j, k, m)] for i in SUBJECTMAPPING[j]) <= maxclasssize

    # Solving the model
    model += (100 * lpSum(studentsum[(i)] for i in STUDENTS) + lpSum(num930classes[(i)] for i in TIMES) + 500 * lpSum(
        daysforteacherssum[(m)] for m in TEACHERS))
    print("Solving Model")
    model.solve()
    print("Status:", LpStatus[model.status])
    print("Complete")
    add_classes_to_timetable(TEACHERS, TEACHERMAPPING, SUBJECTMAPPING, TIMES, subject_vars, assign_vars)
    print("Status:", LpStatus[model.status])
    return LpStatus[model.status]


def runtimetable_with_rooms_two_step(STUDENTS, SUBJECTS, TIMES, day, DAYS, TEACHERS, SUBJECTMAPPING, REPEATS,
                                     TEACHERMAPPING,
                                     TUTORAVAILABILITY, maxclasssize, minclasssize, ROOMS):
    '''
    Run the timetabling process and input into the database.

    This process calls the CBCSolver using the PuLP package and then adds the classes to the database.


    :param STUDENTS: should be an array of student names
    :param SUBJECTS: should be an array of subject codes
    :param TIMES: an array of strings representing possible timeslots
    :param day:
    :param DAYS: the days corresponding to the timeslots above
    :param TEACHERS: an array of the names of the tutors
    :param SUBJECTMAPPING: This is a dictionary representing the subjects
                            each tutor is taking
    :param REPEATS: A dictionary of how many repeats each subject has
    :param TEACHERMAPPING: A dictionary of what subject each tutor teachers
    :param TUTORAVAILABILITY:
    :param maxclasssize: An integer representing the maximum class size
    :param minclasssize: An integer representing the minimum class size
    :param nrooms: An integer representing the max allowable concurrent classes
    :return: A string representing model status.
    '''
    print("Running solver")
    model = LpProblem('Timetabling', LpMinimize)
    # Create Variables
    print("Creating Variables")
    app.logger.info('Assignment Variables')
    assign_vars = LpVariable.dicts("StudentVariables",
                                   [(i, j, k, m) for m in TEACHERS for j in TEACHERMAPPING[m] for i in SUBJECTMAPPING[j]
                                    for k in TIMES], 0, 1, LpBinary)
    app.logger.info('Subject Variables')
    subject_vars = LpVariable.dicts("SubjectVariables",
                                    [(j, k, m) for m in TEACHERS for j in TEACHERMAPPING[m] for k in TIMES], 0, 1,
                                    LpBinary)

    # c
    app.logger.info('9:30 classes')
    num930classes = LpVariable.dicts("930Classes", [(i) for i in TIMES], lowBound=0, cat=LpInteger)
    # w
    app.logger.info('Days for teachers')
    daysforteachers = LpVariable.dicts("numdaysforteachers", [(i, j) for i in TEACHERS for j in range(len(DAYS))], 0, 1,
                                       LpBinary)
    # p
    daysforteacherssum = LpVariable.dicts("numdaysforteacherssum", [(i) for i in TEACHERS], 0, cat=LpInteger)
    # variables for student clashes
    studenttime = LpVariable.dicts("StudentTime", [(i, j) for i in STUDENTS for j in TIMES], lowBound=0, upBound=1,
                                   cat=LpBinary)
    studentsum = LpVariable.dicts("StudentSum", [(i) for i in STUDENTS], 0, cat=LpInteger)

    # Count the days that a teacher is rostered on. Make it bigger than a small number times the sum
    # for that particular day.
    for m in TEACHERS:
        app.logger.info('Counting Teachers for ' + m)
        for d in range(len(day)):
            model += daysforteachers[(m, d)] >= 0.1 * lpSum(
                subject_vars[(j, k, m)] for j in TEACHERMAPPING[m] for k in DAYS[day[d]])
            model += daysforteachers[(m, d)] <= lpSum(
                subject_vars[(j, k, m)] for j in TEACHERMAPPING[m] for k in DAYS[day[d]])
    for m in TEACHERS:
        model += daysforteacherssum[(m)] == lpSum(daysforteachers[(m, d)] for d in range(len(day)))

    print("Constraining tutor availability")
    # This bit of code puts in the constraints for the tutor availability.
    # It reads in the 0-1 matrix of tutor availability and constrains that no classes
    # can be scheduled when a tutor is not available.
    # The last column of the availabilities is the tutor identifying number, hence why we have
    # used a somewhat convoluted idea down here.
    for m in TEACHERS:
        for k in TIMES:
            if k not in TUTORAVAILABILITY[m]:
                model += lpSum(subject_vars[(j, k, m)] for j in TEACHERMAPPING[m]) == 0

    # Constraints on subjects for each students
    print("Constraining student subjects")
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            for i in SUBJECTMAPPING[j]:
                model += lpSum(assign_vars[(i, j, k, m)] for k in TIMES) == 1

    # This code means that students cannot attend a tute when a tute is not running
    # But can not attend a tute if they attend a repeat.
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            for i in SUBJECTMAPPING[j]:
                for k in TIMES:
                    model += assign_vars[(i, j, k, m)] <= subject_vars[(j, k, m)]

    # Constraints on which tutor can take each class
    # This goes through each list and either constrains it to 1 or 0 depending if
    # the teacher needs to teach that particular class.
    print("Constraining tutor classes")
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            model += lpSum(subject_vars[(j, k, m)] for k in TIMES) == REPEATS[j]

    # General Constraints on Rooms etc.
    print("Constraining times")
    # For each time cannot exceed number of rooms
    for k in TIMES:
        model += lpSum(subject_vars[(j, k, m)] for m in TEACHERS for j in TEACHERMAPPING[m]) <= len(ROOMS)

    # Teachers can only teach one class at a time
    for k in TIMES:
        for m in TEACHERS:
            model += lpSum(subject_vars[(j, k, m)] for j in TEACHERMAPPING[m]) <= 1
    print("Constraint: Minimize student clashes")
    # STUDENT CLASHES
    for i in STUDENTS:
        for k in TIMES:
            model += studenttime[(i, k)] <= lpSum(
                assign_vars[(i, j, k, m)] for m in TEACHERS for j in TEACHERMAPPING[m] if i in SUBJECTMAPPING[j]) / 2
            model += studenttime[(i, k)] >= 0.3 * (0.5 * lpSum(
                assign_vars[(i, j, k, m)] for m in TEACHERS for j in TEACHERMAPPING[m] if i in SUBJECTMAPPING[j]) - 0.5)
    for i in STUDENTS:
        model += studentsum[(i)] == lpSum(studenttime[(i, k)] for k in TIMES)

    # This minimizes the number of 9:30 classes.
    for i in TIMES:
        if i.find('21:30') != -1:
            model += num930classes[(i)] == lpSum(subject_vars[(j, i, m)] for m in TEACHERS for j in TEACHERMAPPING[m])

        else:
            model += num930classes[(i)] == 0

    print("Setting objective function")

    # Class size constraint
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            for k in TIMES:
                model += lpSum(assign_vars[(i, j, k, m)] for i in SUBJECTMAPPING[j]) >= minclasssize * subject_vars[
                    (j, k, m)]
                model += lpSum(assign_vars[(i, j, k, m)] for i in SUBJECTMAPPING[j]) <= maxclasssize

    # Solving the model
    model += (100 * lpSum(studentsum[(i)] for i in STUDENTS) + lpSum(num930classes[(i)] for i in TIMES) + 500 * lpSum(
        daysforteacherssum[(m)] for m in TEACHERS))
    print("Solving Model")
    model.solve()
    print("Status:", LpStatus[model.status])
    print("Completed Timetable")

    print("Allocating Rooms")
    model2 = LpProblem('RoomAllocation', LpMinimize)
    print("Defining Variables")
    subject_vars_rooms = LpVariable.dicts("SubjectVariablesRooms",
                                          [(j, k, m, n) for m in TEACHERS for j in TEACHERMAPPING[m] for k in TIMES for
                                           n in ROOMS], 0, 1, LpBinary)

    teacher_number_rooms = LpVariable.dicts("NumberRoomsTeacher", [(m, n) for m in TEACHERS for n in ROOMS], 0, 1,
                                            LpBinary)
    teacher_number_rooms_sum = LpVariable.dicts("NumberRoomsTeacherSum", [(m) for m in TEACHERS], 0, cat=LpInteger)

    print("Minimizing number of rooms for each tutor")
    for m in TEACHERS:
        for n in ROOMS:
            model2 += teacher_number_rooms[(m, n)] >= 0.01 * lpSum(
                subject_vars_rooms[(j, k, m, n)] for j in TEACHERMAPPING[m] for k in TIMES)
            model2 += teacher_number_rooms[(m, n)] <= lpSum(
                subject_vars_rooms[(j, k, m, n)] for j in TEACHERMAPPING[m] for k in TIMES)
    print("1")
    for m in TEACHERS:
        model2 += teacher_number_rooms_sum[(m)] == lpSum(teacher_number_rooms[(m, n)] for n in ROOMS)

    # Rooms must be allocated at times when the classes are running
    print("Constraining Times")
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            for k in TIMES:
                model2 += lpSum(subject_vars_rooms[(j, k, m, n)] for n in ROOMS) == subject_vars[(j, k, m)].varValue

    print("Ensuring Uniqueness")
    # Can only have one class in each room at a time.
    for k in TIMES:
        for n in ROOMS:
            model2 += lpSum(subject_vars_rooms[(j, k, m, n)] for m in TEACHERS for j in TEACHERMAPPING[m]) <= 1

    print("Setting Objective Function")
    model2 += lpSum(teacher_number_rooms_sum[(m)] for m in TEACHERS)
    print("Solve Room Allocation")
    model2.solve()
    print("Complete")
    print("Adding to Database")
    Attendance.models.add_classes_to_timetable_twostep(TEACHERS, TEACHERMAPPING, SUBJECTMAPPING, TIMES,
                                                       subject_vars_rooms, assign_vars, ROOMS)
    print("Status:", LpStatus[model.status])
    return LpStatus[model.status]


def runtimetable_with_rooms(STUDENTS, SUBJECTS, TIMES, day, DAYS, TEACHERS, SUBJECTMAPPING, REPEATS, TEACHERMAPPING,
                            TUTORAVAILABILITY, maxclasssize, minclasssize, ROOMS):
    '''
    Run the timetabling process and input into the database.

    This process calls the CBCSolver using the PuLP package and then adds the classes to the database.


    :param STUDENTS: should be an array of student names
    :param SUBJECTS: should be an array of subject codes
    :param TIMES: an array of strings representing possible timeslots
    :param day:
    :param DAYS: the days corresponding to the timeslots above
    :param TEACHERS: an array of the names of the tutors
    :param SUBJECTMAPPING: This is a dictionary representing the subjects
                            each tutor is taking
    :param REPEATS: A dictionary of how many repeats each subject has
    :param TEACHERMAPPING: A dictionary of what subject each tutor teachers
    :param TUTORAVAILABILITY:
    :param maxclasssize: An integer representing the maximum class size
    :param minclasssize: An integer representing the minimum class size
    :param nrooms: An integer representing the max allowable concurrent classes
    :return: A string representing model status.
    '''
    print("Running solver")

    model = LpProblem('Timetabling', LpMinimize)
    # Create Variables
    print("Creating Variables")
    print('Assignment Variables')
    assign_vars = LpVariable.dicts("StudentVariables",
                                   [(i, j, k, m, n) for m in TEACHERS for j in TEACHERMAPPING[m] for i in
                                    SUBJECTMAPPING[j]
                                    for k in TIMES for n in ROOMS], 0, 1, LpBinary)
    app.logger.info('Subject Variables')
    subject_vars = LpVariable.dicts("SubjectVariables",
                                    [(j, k, m, n) for m in TEACHERS for j in TEACHERMAPPING[m] for k in TIMES for n in
                                     ROOMS], 0, 1,
                                    LpBinary)

    # c
    app.logger.info('9:30 classes')
    num930classes = LpVariable.dicts("930Classes", [(i) for i in TIMES], lowBound=0, cat=LpInteger)
    # w
    print('Days for teachers')
    daysforteachers = LpVariable.dicts("numdaysforteachers", [(i, j) for i in TEACHERS for j in range(len(DAYS))], 0, 1,
                                       LpBinary)
    # p
    daysforteacherssum = LpVariable.dicts("numdaysforteacherssum", [(i) for i in TEACHERS], 0, cat=LpInteger)
    # variables for student clashes
    studenttime = LpVariable.dicts("StudentTime", [(i, j) for i in STUDENTS for j in TIMES], lowBound=0, upBound=1,
                                   cat=LpBinary)
    studentsum = LpVariable.dicts("StudentSum", [(i) for i in STUDENTS], 0, cat=LpInteger)

    # Count the days that a teacher is rostered on. Make it bigger than a small number times the sum
    # for that particular day.
    for m in TEACHERS:
        print('Counting Teachers for ' + m)
        for d in range(len(day)):
            model += daysforteachers[(m, d)] >= 0.1 * lpSum(
                subject_vars[(j, k, m, n)] for j in TEACHERMAPPING[m] for k in DAYS[day[d]] for n in ROOMS)
            model += daysforteachers[(m, d)] <= lpSum(
                subject_vars[(j, k, m, n)] for j in TEACHERMAPPING[m] for k in DAYS[day[d]] for n in ROOMS)
    for m in TEACHERS:
        model += daysforteacherssum[(m)] == lpSum(daysforteachers[(m, d)] for d in range(len(day)))

    print("Constraining tutor availability")
    # This bit of code puts in the constraints for the tutor availability.
    # It reads in the 0-1 matrix of tutor availability and constrains that no classes
    # can be scheduled when a tutor is not available.
    # The last column of the availabilities is the tutor identifying number, hence why we have
    # used a somewhat convoluted idea down here.
    for m in TEACHERS:
        for k in TIMES:
            if k not in TUTORAVAILABILITY[m]:
                model += lpSum(subject_vars[(j, k, m, n)] for j in TEACHERMAPPING[m] for n in ROOMS) == 0

    # Constraints on subjects for each students
    print("Constraining student subjects")
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            for i in SUBJECTMAPPING[j]:
                model += lpSum(assign_vars[(i, j, k, m, n)] for k in TIMES for n in ROOMS) == 1

    # There can only be one class in each room at a time
    print("Constraining Rooms")
    for n in ROOMS:
        for k in TIMES:
            model += lpSum(subject_vars[(j, k, m, n)] for m in TEACHERS for j in TEACHERMAPPING[m]) <= 1

    # This code means that students cannot attend a tute when a tute is not running
    # But can not attend a tute if they attend a repeat.
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            for i in SUBJECTMAPPING[j]:
                for k in TIMES:
                    for n in ROOMS:
                        model += assign_vars[(i, j, k, m, n)] <= subject_vars[(j, k, m, n)]

    # Constraints on which tutor can take each class
    # This goes through each list and either constrains it to 1 or 0 depending if
    # the teacher needs to teach that particular class.
    print("Constraining tutor classes")
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            model += lpSum(subject_vars[(j, k, m, n)] for k in TIMES for n in ROOMS) == REPEATS[j]

    # General Constraints on Rooms etc.
    print("Constraining times")
    # For each time cannot exceed number of rooms
    for k in TIMES:
        model += lpSum(subject_vars[(j, k, m, n)] for m in TEACHERS for j in TEACHERMAPPING[m] for n in ROOMS) <= len(
            ROOMS)

    # Teachers can only teach one class at a time
    for k in TIMES:
        for m in TEACHERS:
            model += lpSum(subject_vars[(j, k, m, n)] for j in TEACHERMAPPING[m] for n in ROOMS) <= 1
    print("Constraint: Minimize student clashes")
    # STUDENT CLASHES
    for i in STUDENTS:
        for k in TIMES:
            model += studenttime[(i, k)] <= lpSum(
                assign_vars[(i, j, k, m, n)] for m in TEACHERS for j in TEACHERMAPPING[m] for n in ROOMS if
                i in SUBJECTMAPPING[j]) / 2
            model += studenttime[(i, k)] >= 0.3 * (0.5 * lpSum(
                assign_vars[(i, j, k, m, n)] for m in TEACHERS for j in TEACHERMAPPING[m] for n in ROOMS if
                i in SUBJECTMAPPING[j]) - 0.5)
    for i in STUDENTS:
        model += studentsum[(i)] == lpSum(studenttime[(i, k)] for k in TIMES)

    # This minimizes the number of 9:30 classes.
    for i in TIMES:
        if i.find('21:30') != -1:
            model += num930classes[(i)] == lpSum(
                subject_vars[(j, i, m, n)] for m in TEACHERS for j in TEACHERMAPPING[m] for n in ROOMS)

        else:
            model += num930classes[(i)] == 0

    print("Class Size Constraint")
    # Class size constraint
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            for k in TIMES:
                for n in ROOMS:
                    model += lpSum(assign_vars[(i, j, k, m, n)] for i in SUBJECTMAPPING[j]) >= minclasssize * \
                                                                                               subject_vars[
                                                                                                   (j, k, m, n)]
                    model += lpSum(assign_vars[(i, j, k, m, n)] for i in SUBJECTMAPPING[j]) <= maxclasssize

    # Solving the model
    print("Setting objective function")
    model += (100 * lpSum(studentsum[(i)] for i in STUDENTS) + lpSum(num930classes[(i)] for i in TIMES) + 500 * lpSum(
        daysforteacherssum[(m)] for m in TEACHERS))
    print("Solving Model")
    model.solve()
    print("Status:", LpStatus[model.status])
    print("Complete")
    add_classes_to_timetable(TEACHERS, TEACHERMAPPING, SUBJECTMAPPING, TIMES, subject_vars, assign_vars, ROOMS)
    print("Status:", LpStatus[model.status])
    return LpStatus[model.status]

def preparetimetable(addtonewtimetable=False):
    '''
    Get timetable data and then execute the timetabling program.

    :param addtonewtimetable: Whether this should be added to a new timetable and set as default.
    :return: The view timetable page.
    '''
    # if addtonewtimetable == "true":
    #    timetable = Timetable(get_current_year(),get_current_studyperiod())
    #    db.session.add(timetable)
    #    db.session.commit()
    #    admin = Admin.query.filter_by(key="timetable").first()
    #    admin.value = timetable.id
    #    db.session.commit()
    print("Preparing Timetable")
    # (STUDENTS, SUBJECTS, TIMES, day, DAYS, TEACHERS, SUBJECTMAPPING, REPEATS, TEACHERMAPPING,
    #TUTORAVAILABILITY, maxclasssize, minclasssize, nrooms) = Attendance.models.get_timetable_data()

    (STUDENTS, SUBJECTS, TIMES, day, DAYS, TEACHERS, SUBJECTMAPPING, REPEATS, TEACHERMAPPING,
     TUTORAVAILABILITY, maxclasssize, minclasssize, ROOMS) = Attendance.models.get_timetable_data(rooms=True)
    print("Everything ready")
    executor.submit(runtimetable_with_rooms_two_step, STUDENTS, SUBJECTS, TIMES, day, DAYS, TEACHERS, SUBJECTMAPPING,
                    REPEATS, TEACHERMAPPING,
                    TUTORAVAILABILITY, maxclasssize, minclasssize, ROOMS)
    form = AddTimetableForm()
    return render_template("viewtimetable.html", form=form)




def allowed_file(filename):
    '''
    Checks whether the uploaded file has an allowed extension.
    :param filename: The filename to check
    :return: True/False
    '''
    return '.' in filename and \
           filename.rsplit('.', 1)[1] in app.config['ALLOWED_EXTENSIONS']


def upload(file):
    '''
    Save the uploaded file to the UPLOAD_FOLDER directory.

    :param file: The file to upload
    :return: Filename of uploaded file in upload folder
    '''
    if file and allowed_file(file.filename):
        # Make the filename safe, remove unsupported chars
        filename = file.filename
        # Move the file form the temporal folder to
        # the upload folder we setup
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        filename2 = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        return filename2


def checkboxvalue(checkbox):
    '''
    Get value of checkbox.

    :param checkbox: Input from request.form
    :return: 1 if ticked, 0 if not.
    '''
    if (checkbox != None):
        return 1
    else:
        return 0


def read_excel(filename):
    '''
    Read Excel File provided by filename.

    :param filename - path to an Excel file:
    :return: pandas dataframe
    '''
    xl = ExcelFile(filename)
    df = xl.parse(xl.sheet_names[0])
    return df


def create_roll(students, subject, timeslot, room):
    document = Document()

    document.add_heading(subject.subname, 0)

    document.add_paragraph('Timeslot: ' + timeslot.day + " " + timeslot.time)
    if room is not None:
        document.add_paragraph('Room: ' + room.name)

    table = document.add_table(rows=1, cols=12)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = '1'
    hdr_cells[2].text = '2'
    hdr_cells[3].text = '3'
    hdr_cells[4].text = '4'
    hdr_cells[5].text = '5'
    hdr_cells[6].text = '6'
    hdr_cells[7].text = '7'
    hdr_cells[8].text = '8'
    hdr_cells[9].text = '9'
    hdr_cells[10].text = '10'
    hdr_cells[11].text = '11'
    for item in students:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.name)
    document.save(app.config['UPLOAD_FOLDER'] + '/' + subject.subcode + '.docx')
    return app.config['UPLOAD_FOLDER'] + '/' + subject.subcode + '.docx'
    # document.save('demo.docx')


def create_excel(data):
    writer = pandas.ExcelWriter(app.config['UPLOAD_FOLDER'] + '/timetable.xlsx', engine='xlsxwriter')
    data.to_excel(writer, sheet_name='Timetable', index=False)
    writer.save()
    return app.config['UPLOAD_FOLDER'] + '/timetable.xlsx'


def format_timetable_data_for_export():
    timeslots = Attendance.models.get_all_timeslots()
    timeslots = sorted(timeslots, key=attrgetter('daynumeric', 'time'))
    timetable = []
    for i in range(len(timeslots)):
        timeslot = timeslots[i]
        classes = timeslot.timetabledclasses
        for timeclass in classes:
            if timeclass.room is not None:
                room = timeclass.room.name
            else:
                room = ""
            if timeclass.tutor is not None:
                tutor = timeclass.tutor.name
            else:
                tutor = ""

            timetable.append((timeclass.timeslot.day + ' ' + timeclass.timeslot.time, timeclass.subject.subname,
                              tutor, room))

    timetable = pandas.DataFrame(timetable)
    timetable.columns = ['Time', 'Subject', 'Tutor', 'Room']
    return timetable




def format_tutor_hours_for_export(hours):
    hours = list(hours)
    hours = pandas.DataFrame(hours)
    hours.columns = ['Name', 'Initial Tutorials', 'Repeat Tutorials']
    return hours
